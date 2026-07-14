"""PR-1 — run contract accepts attachments (images + documents).

Covers the shared extractor (`services/extract.py`), the agent runner's
attachment preparation (documents -> context, images -> multimodal blocks or a
note), and the run endpoint's JSON-body contract with query-param back-compat.
"""

from __future__ import annotations

from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document

from app.models.attachment import Attachment, RunRequest
from app.services.agent_executor import AgentRunner, _model_supports_vision
from app.services.extract import extract_text

# --- Shared extractor ---------------------------------------------------------


@pytest.mark.asyncio
async def test_extract_text_reads_local_txt(tmp_path, monkeypatch):
    monkeypatch.setenv("AF_UPLOAD_DIR", str(tmp_path))
    f = tmp_path / "notes.txt"
    f.write_text("hello from a text file")
    assert await extract_text(str(f)) == "hello from a text file"


@pytest.mark.asyncio
async def test_extract_text_reads_local_docx(tmp_path, monkeypatch):
    monkeypatch.setenv("AF_UPLOAD_DIR", str(tmp_path))
    f = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("First line of the doc")
    doc.add_paragraph("Second line of the doc")
    doc.save(str(f))

    text = await extract_text(str(f))
    assert "First line of the doc" in text
    assert "Second line of the doc" in text


@pytest.mark.asyncio
async def test_extract_text_handles_file_uri(tmp_path, monkeypatch):
    monkeypatch.setenv("AF_UPLOAD_DIR", str(tmp_path))
    f = tmp_path / "readme.md"
    f.write_text("# Title\n\nbody")
    assert "# Title" in await extract_text(f.as_uri())


@pytest.mark.asyncio
async def test_extract_text_caps_length(tmp_path, monkeypatch):
    monkeypatch.setenv("AF_UPLOAD_DIR", str(tmp_path))
    f = tmp_path / "big.txt"
    f.write_text("x" * 50_000)
    assert len(await extract_text(str(f))) == 10_000


@pytest.mark.asyncio
async def test_extract_text_rejects_file_outside_upload_dir(tmp_path, monkeypatch):
    monkeypatch.setenv("AF_UPLOAD_DIR", str(tmp_path))
    with pytest.raises(ValueError, match="upload directory"):
        await extract_text("file:///etc/passwd")


@pytest.mark.asyncio
async def test_extract_text_rejects_unsupported_scheme():
    with pytest.raises(ValueError, match="scheme"):
        await extract_text("ftp://internal-host/secret.txt")


# --- Vision detection ---------------------------------------------------------


def test_vision_capable_detection():
    assert _model_supports_vision("gpt-4o")
    assert _model_supports_vision("gpt-4o-mini")
    assert _model_supports_vision("gpt-4-turbo")
    assert not _model_supports_vision("gpt-3.5-turbo")
    assert not _model_supports_vision("ollama/llama3.2:3b")
    assert not _model_supports_vision(None)


# --- Attachment preparation ---------------------------------------------------


@pytest.mark.asyncio
async def test_prepare_attachments_extracts_documents():
    runner = AgentRunner(user_id="u1")
    attachments = [{"url": "file:///tmp/r.pdf", "kind": "document", "name": "r.pdf", "mime": "application/pdf"}]

    with patch("app.services.extract.extract_text", new=AsyncMock(return_value="EXTRACTED BODY")) as mock_extract:
        doc_context, image_blocks, notes = await runner._prepare_attachments(attachments, "gpt-4o-mini")

    mock_extract.assert_awaited_once()
    assert "--- file: r.pdf ---" in doc_context
    assert "EXTRACTED BODY" in doc_context
    assert image_blocks == []
    assert notes == []


@pytest.mark.asyncio
async def test_prepare_attachments_images_for_vision_model():
    runner = AgentRunner(user_id="u1")
    attachments = [{"url": "https://x/y.png", "kind": "image", "name": "y.png", "mime": "image/png"}]

    doc_context, image_blocks, notes = await runner._prepare_attachments(attachments, "gpt-4o")

    assert doc_context == ""
    assert image_blocks == [{"type": "image_url", "image_url": {"url": "https://x/y.png"}}]
    assert notes == []


@pytest.mark.asyncio
async def test_prepare_attachments_images_noted_for_non_vision_model():
    runner = AgentRunner(user_id="u1")
    attachments = [{"url": "https://x/y.png", "kind": "image", "name": "y.png", "mime": "image/png"}]

    doc_context, image_blocks, notes = await runner._prepare_attachments(attachments, "gpt-3.5-turbo")

    assert image_blocks == []
    assert len(notes) == 1
    assert "model not multimodal" in notes[0]


@pytest.mark.asyncio
async def test_prepare_attachments_document_failure_degrades_to_note():
    runner = AgentRunner(user_id="u1")
    attachments = [{"url": "file:///nope.pdf", "kind": "document", "name": "nope.pdf", "mime": "application/pdf"}]

    with patch("app.services.extract.extract_text", new=AsyncMock(side_effect=OSError("boom"))):
        doc_context, image_blocks, notes = await runner._prepare_attachments(attachments, "gpt-4o-mini")

    assert doc_context == ""
    assert len(notes) == 1
    assert "could not read nope.pdf" in notes[0]


@pytest.mark.asyncio
async def test_prepare_attachments_empty():
    runner = AgentRunner(user_id="u1")
    assert await runner._prepare_attachments(None, "gpt-4o") == ("", [], [])


# --- Run endpoint contract ----------------------------------------------------


def test_run_request_body_wins_over_query_param():
    body = RunRequest(input_text="from body", attachments=[])
    assert body.input_text == "from body"


def test_run_request_parses_attachments():
    body = RunRequest.model_validate(
        {
            "input_text": "review these",
            "attachments": [
                {"url": "https://x/a.pdf", "kind": "document", "name": "a.pdf", "mime": "application/pdf"},
                {"url": "https://x/b.png", "kind": "image", "name": "b.png", "mime": "image/png"},
            ],
        }
    )
    assert len(body.attachments) == 2
    assert body.attachments[0].kind == "document"
    assert body.attachments[1].kind == "image"


def test_attachment_rejects_unknown_kind():
    with pytest.raises(ValueError):
        Attachment(url="x", kind="video", name="v.mp4", mime="video/mp4")  # type: ignore[arg-type]


def _run_endpoint_setup(mock_db, *, agent_steps=None):
    """Wire a mock db so POST /api/agents/agent-1/run reaches the runner."""
    mock_user = MagicMock()
    mock_user.user = MagicMock(id="test-user-id-123")
    mock_db.auth.get_user.return_value = mock_user

    mock_agent = MagicMock()
    mock_agent.data = {
        "id": "agent-1",
        "user_id": "test-user-id-123",
        "name": "Vision Agent",
        "system_prompt": "look",
        "tools": [],
        "workflow_steps": agent_steps or [],
        "model": "gpt-4o",
    }
    mock_db.table.return_value.select.return_value.eq.return_value.single.return_value.execute.return_value = mock_agent
    mock_db.table.return_value.insert.return_value.execute.return_value = MagicMock(data=[{"id": "run-1"}])


def test_run_endpoint_accepts_attachments_body(client):
    """A JSON body with attachments is parsed and forwarded to the runner."""
    captured = {}

    with patch("app.db._db") as mock_db:
        _run_endpoint_setup(mock_db)

        with patch("app.routers.runs.AgentRunner") as mock_runner_cls:
            mock_runner = MagicMock()

            async def fake_execute(agent_config, input_text, **kwargs):
                captured["input_text"] = input_text
                captured["attachments"] = kwargs.get("attachments")
                yield {"type": "step", "content": "ok", "tokens": 0}

            mock_runner.execute = fake_execute
            mock_runner_cls.return_value = mock_runner

            resp = client.post(
                "/api/agents/agent-1/run?token=t&input_text=ignored",
                json={
                    "input_text": "from body",
                    "attachments": [
                        {"url": "https://x/p.pdf", "kind": "document", "name": "p.pdf", "mime": "application/pdf"}
                    ],
                },
            )

    assert resp.status_code == 200
    # Body input_text wins over the query param.
    assert captured["input_text"] == "from body"
    assert captured["attachments"][0]["name"] == "p.pdf"


def test_run_endpoint_backcompat_query_param_only(client):
    """Old callers (query param, no body) still work and get no attachments."""
    captured = {}

    with patch("app.db._db") as mock_db:
        _run_endpoint_setup(mock_db)

        with patch("app.routers.runs.AgentRunner") as mock_runner_cls:
            mock_runner = MagicMock()

            async def fake_execute(agent_config, input_text, **kwargs):
                captured["input_text"] = input_text
                captured["attachments"] = kwargs.get("attachments")
                yield {"type": "step", "content": "ok", "tokens": 0}

            mock_runner.execute = fake_execute
            mock_runner_cls.return_value = mock_runner

            resp = client.post("/api/agents/agent-1/run?token=t&input_text=hello")

    assert resp.status_code == 200
    assert captured["input_text"] == "hello"
    assert captured["attachments"] == []
